"""
SID DOCX Export
===============
Renders the SID dashboard slide spec as a formatted Word document with
two tables matching the two-slide layout:
  Table 1: Affected suppliers overview (from Slide 1)
  Table 2: Supplier fulfillment detail (from Slide 2)

Uses python-docx for generation, consistent with gfd_docx_export.py.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Colour palette ─────────────────────────────────────────────────

_CLR = {
    "HDR_BLUE":   "3F7AB6",
    "DARK_BLUE":  "00457E",
    "GREEN":      "38DF12",
    "YELLOW":     "FCEF39",
    "RED":        "C1001F",
    "BLACK":      "000000",
    "WHITE":      "FFFFFF",
    "GREY":       "A0A0A0",
    "LIGHT_GREY": "F0F0F0",
}

FONT = "Arial"


# ─── Low-level helpers ──────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    """Set solid background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, *,
                   size: float = 8, bold: bool = False,
                   color: str | None = None,
                   align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Write text into a cell with font formatting."""
    cell.text = ""
    lines = str(text).split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(size + 2)
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _set_row_height(row_obj, val: int = 300, rule: str = "atLeast") -> None:
    """Set row height via XML."""
    tr = row_obj._tr
    trPr = tr.get_or_add_trPr()
    for existing in trPr.findall(qn("w:trHeight")):
        trPr.remove(existing)
    h = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{val}" w:hRule="{rule}"/>')
    trPr.append(h)


def _set_table_borders(table, color: str = "AAAAAA", size: int = 4) -> None:
    """Set thin borders on the whole table."""
    tbl_xml = table._tbl
    tblPr = tbl_xml.tblPr if tbl_xml.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>',
    )
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))


# ─── Main export function ───────────────────────────────────────────

def sid_spec_to_docx(slide_spec: dict) -> io.BytesIO:
    """
    Convert a SID slide spec dict into a formatted Word document.

    Parameters
    ----------
    slide_spec : dict produced by llm_generate_slide_spec or fallback

    Returns
    -------
    io.BytesIO buffer (seeked to 0) containing the .docx bytes
    """
    doc = Document()

    # ── Page setup: landscape ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # ── Title ──
    title_text = slide_spec.get("presentation_title", "Supplier Situation Update")
    last_update = slide_spec.get("last_update", "")
    if last_update:
        title_text = f"{title_text}  —  Last Update: {last_update}"

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_after = Pt(4)
    run = title_para.add_run(title_text)
    run.font.name = FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(_CLR["DARK_BLUE"])

    # ── Evaluation summary ──
    eval_summary = slide_spec.get("evaluation_summary", "")
    if eval_summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(eval_summary)
        run.font.name = FONT
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(_CLR["DARK_BLUE"])

    # ── Coverage Distribution ──
    cov = slide_spec.get("coverage_distribution", {})
    if any(cov.values()):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("Coverage Distribution: ")
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.bold = True
        items = [
            f"No coverage: {cov.get('no_coverage', 0)}",
            f"< 4 days: {cov.get('lt_4_days', 0)}",
            f"5–15 days: {cov.get('5_to_15_days', 0)}",
            f"> 15 days: {cov.get('gt_15_days', 0)}",
        ]
        run2 = p.add_run("  |  ".join(items))
        run2.font.name = FONT
        run2.font.size = Pt(9)

    # ── Section 1: Affected Suppliers Table ──
    affected = slide_spec.get("affected_suppliers", [])
    if affected:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run("Affected Suppliers")
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(_CLR["DARK_BLUE"])

        n_rows = len(affected) + 1
        table1 = doc.add_table(rows=n_rows, cols=5)
        table1.alignment = WD_TABLE_ALIGNMENT.LEFT
        table1.autofit = False

        col_widths = [Inches(2.0), Inches(0.7), Inches(0.7), Inches(0.7), Inches(6.0)]
        for ci, w in enumerate(col_widths):
            for row in table1.rows:
                row.cells[ci].width = w

        # Header
        headers = ["Supplier Name", "Cat", "Q-PAVE", "L-PAVE", "Remarks"]
        for ci, label in enumerate(headers):
            cell = table1.cell(0, ci)
            _set_cell_shading(cell, _CLR["HDR_BLUE"])
            _set_cell_text(cell, label, size=8, bold=True, color=_CLR["WHITE"],
                           align=WD_ALIGN_PARAGRAPH.CENTER if ci in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
        _set_row_height(table1.rows[0], 340)

        # Data
        for ri, sup in enumerate(affected):
            row_idx = ri + 1
            _set_cell_text(table1.cell(row_idx, 0), sup.get("supplier_name", ""), size=8)
            _set_cell_text(table1.cell(row_idx, 1), sup.get("cat", ""), size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(table1.cell(row_idx, 2), str(sup.get("q_pave", "")), size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(table1.cell(row_idx, 3), str(sup.get("l_pave", "")), size=8,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(table1.cell(row_idx, 4), sup.get("remarks", ""), size=7)
            _set_row_height(table1.rows[row_idx], 260)

        _set_table_borders(table1)

    # ── Section 2: Actions Table ──
    actions = slide_spec.get("actions", [])
    if actions:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run("Actions Moving Forward")
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(_CLR["DARK_BLUE"])

        n_rows = len(actions) + 1
        table2 = doc.add_table(rows=n_rows, cols=4)
        table2.alignment = WD_TABLE_ALIGNMENT.LEFT
        table2.autofit = False

        col_widths = [Inches(2.5), Inches(1.5), Inches(1.0), Inches(5.1)]
        for ci, w in enumerate(col_widths):
            for row in table2.rows:
                row.cells[ci].width = w

        act_headers = ["Action", "Resp.", "Deadline", "Status / Comments"]
        for ci, label in enumerate(act_headers):
            cell = table2.cell(0, ci)
            _set_cell_shading(cell, _CLR["HDR_BLUE"])
            _set_cell_text(cell, label, size=8, bold=True, color=_CLR["WHITE"])
        _set_row_height(table2.rows[0], 340)

        for ri, act in enumerate(actions):
            row_idx = ri + 1
            _set_cell_text(table2.cell(row_idx, 0), act.get("action", ""), size=8)
            _set_cell_text(table2.cell(row_idx, 1), act.get("resp", ""), size=8)
            _set_cell_text(table2.cell(row_idx, 2), act.get("deadline", ""), size=8)
            _set_cell_text(table2.cell(row_idx, 3), act.get("status_comments", ""), size=7)
            _set_row_height(table2.rows[row_idx], 260)

        _set_table_borders(table2)

    # ── Contextual Notes ──
    notes = slide_spec.get("contextual_notes", "")
    if notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(notes)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(_CLR["GREY"])

    # ── Page break before detail table ──
    doc.add_page_break()

    # ── Section 3: Supplier Fulfillment Detail ──
    details = slide_spec.get("supplier_details", [])
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run("Fulfillment Overview (impacted suppliers)")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(_CLR["DARK_BLUE"])

    if details:
        n_rows = len(details) + 1
        table3 = doc.add_table(rows=n_rows, cols=9)
        table3.alignment = WD_TABLE_ALIGNMENT.LEFT
        table3.autofit = False

        col_widths = [
            Inches(1.2), Inches(0.9), Inches(0.9), Inches(0.8),
            Inches(1.0), Inches(1.1), Inches(1.0), Inches(1.0), Inches(2.2),
        ]
        for ci, w in enumerate(col_widths):
            for row in table3.rows:
                row.cells[ci].width = w

        detail_headers = [
            "Supplier\nname", "Host", "Material\nPlanner", "SDA",
            "Coverage\nDate", "Coverage\nafter actions",
            "Affected\nproduct", "Customer", "Remarks",
        ]
        for ci, label in enumerate(detail_headers):
            cell = table3.cell(0, ci)
            _set_cell_shading(cell, _CLR["DARK_BLUE"])
            _set_cell_text(cell, label, size=7, bold=True, color=_CLR["WHITE"])
        _set_row_height(table3.rows[0], 400)

        detail_fields = [
            "supplier_name", "host", "material_planner", "sda",
            "coverage_date", "coverage_after_actions",
            "affected_product", "customer", "remarks",
        ]
        for ri, row_data in enumerate(details):
            row_idx = ri + 1
            for ci, field in enumerate(detail_fields):
                text = row_data.get(field, "") or ""
                _set_cell_text(table3.cell(row_idx, ci), str(text), size=7)
            _set_row_height(table3.rows[row_idx], 300)

            # Alternate row shading
            if ri % 2 == 1:
                for ci in range(9):
                    _set_cell_shading(table3.cell(row_idx, ci), _CLR["LIGHT_GREY"])

        _set_table_borders(table3)

    # ── Footer ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_before = Pt(4)
    run = footer_para.add_run(f"Generated {date.today().isoformat()}")
    run.font.name = FONT
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string(_CLR["GREY"])

    # ── Save to buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

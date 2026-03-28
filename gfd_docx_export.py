"""
GFD DOCX Export
===============
Renders the GFD dashboard slide spec as a formatted Word document table.
The user can download this as an alternative to the .pptx and copy-paste
the table directly into PowerPoint or other tools.

The table mirrors the corporate PPTX template:
  14-column layout · 2-row merged header · vertically merged product-group
  cells with amber fill · RAG-coloured CW cells · Arial 6/7pt body/header

Uses python-docx for generation, consistent with the existing docx_export.py.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Colour palette (matching the corporate GFD template) ───────────

_CLR = {
    "HDR_BLUE":  "00457E",
    "HDR_GREEN": "008000",
    "GREEN":     "00B050",
    "AMBER":     "FFC000",
    "RED":       "FF0000",
    "GREY":      "A0A0A0",
    "WHITE":     "FFFFFF",
    "BLACK":     "000000",
}

_RAG_HEX = {
    "GREEN": _CLR["GREEN"],
    "AMBER": _CLR["AMBER"],
    "RED":   _CLR["RED"],
    "GREY":  _CLR["GREY"],
}

FONT = "Arial"

# Column widths in inches for landscape A4 (~11.69 usable minus margins)
# Proportioned from the PPTX template EMU values, scaled to fit landscape page.
_COL_WIDTHS_IN = [
    1.05,   # col 0:  Product Group (PG)
    0.40,   # col 1:  Plant
    1.75,   # col 2:  Customer / Channel
    0.70,   # col 3:  KB Coverage
    0.28,   # col 4:  CW+0
    0.28,   # col 5:  CW+1
    0.28,   # col 6:  CW+2
    0.28,   # col 7:  CW+3
    0.28,   # col 8:  CW+4
    0.28,   # col 9:  CW+5
    0.28,   # col 10: Quarter
    1.90,   # col 11: Supplier
    1.40,   # col 12: Comment
    0.84,   # col 13: FM Detail Letter
]


# ─── Low-level helpers ──────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str) -> None:
    """Set solid background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, *,
                   size: float = 6, bold: bool = False,
                   color: str | None = None,
                   align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Write text into a cell with font formatting. Handles newlines."""
    # Clear default empty paragraph
    cell.text = ""
    lines = str(text).split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(size + 2)
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _set_vmerge_restart(cell) -> None:
    """Mark a cell as the start of a vertical merge."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing vMerge
    for existing in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(existing)
    merge_elem = parse_xml(f'<w:vMerge {nsdecls("w")} w:val="restart"/>')
    tcPr.append(merge_elem)


def _set_vmerge_continue(cell) -> None:
    """Mark a cell as a continuation of a vertical merge."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(existing)
    merge_elem = parse_xml(f'<w:vMerge {nsdecls("w")}/>')
    tcPr.append(merge_elem)


def _set_hmerge_span(table, row_idx: int, start_col: int, end_col: int) -> None:
    """
    Merge cells horizontally from start_col to end_col (inclusive) in a row.
    Uses the python-docx merge API.
    """
    table.cell(row_idx, start_col).merge(table.cell(row_idx, end_col))


def _compute_product_group_ranges(rows: list[dict]) -> list[tuple[int, int, str]]:
    """
    Identify consecutive row spans that share the same product_group value.
    Returns list of (start_idx, end_idx, group_label) — 0-based, inclusive.
    """
    if not rows:
        return []
    ranges: list[tuple[int, int, str]] = []
    current = rows[0].get("product_group", "")
    start = 0
    for i in range(1, len(rows)):
        grp = rows[i].get("product_group", "")
        if grp != current:
            ranges.append((start, i - 1, current))
            current = grp
            start = i
    ranges.append((start, len(rows) - 1, current))
    return ranges


# ─── Main export function ───────────────────────────────────────────

def gfd_spec_to_docx(slide_spec: dict) -> io.BytesIO:
    """
    Convert a GFD slide spec dict into a formatted Word document
    with a colour-coded dashboard table.

    Parameters
    ----------
    slide_spec : dict produced by llm_generate_slide_spec or fallback

    Returns
    -------
    io.BytesIO buffer (seeked to 0) containing the .docx bytes
    """
    doc = Document()

    # ── Page setup: landscape ────────────────────────────────────────
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # A4 landscape
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin    = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)

    # ── Title ────────────────────────────────────────────────────────
    title_text = slide_spec.get("presentation_title", "Global Fulfilment Dashboard")
    cw_label = slide_spec.get("current_cw", "")
    if cw_label:
        title_text = f"{title_text}  —  {cw_label}"

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_after = Pt(6)
    run = title_para.add_run(title_text)
    run.font.name = FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(_CLR["HDR_BLUE"])

    # ── Build the table ──────────────────────────────────────────────
    all_rows: list[dict]  = slide_spec.get("rows", [])
    cw_columns: list[int] = slide_spec.get("cw_columns", [])
    quarter_label: str    = slide_spec.get("quarter_label", "Q?")

    n_data    = len(all_rows)
    n_tbl_rows = n_data + 2     # 2 header rows
    n_cols    = len(_COL_WIDTHS_IN)

    table = doc.add_table(rows=n_tbl_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Column widths
    for ci, w_in in enumerate(_COL_WIDTHS_IN):
        for row in table.rows:
            row.cells[ci].width = Inches(w_in)

    # ── Header row 0 ────────────────────────────────────────────────
    hdr_left = [
        "Product\nGroup\n(PG)",
        "Plant",
        "Customer/\nChannel (PG)",
        "KB Coverage\nto Customer\n(CW/YY)",
    ]
    hdr_right = ["Supplier", "Comment", "FM Detail Letter"]

    CW_START = 4
    Q_COL    = 10

    # Left text headers (dark blue bg, white text) — will be vertically merged
    for ci, label in enumerate(hdr_left):
        cell = table.cell(0, ci)
        _set_cell_shading(cell, _CLR["HDR_BLUE"])
        _set_cell_text(cell, label, size=7, bold=True, color=_CLR["WHITE"])
        _set_vmerge_restart(cell)
        # Row 1 continuation
        cell1 = table.cell(1, ci)
        _set_cell_shading(cell1, _CLR["HDR_BLUE"])
        _set_cell_text(cell1, "", size=7)
        _set_vmerge_continue(cell1)

    # CW band header (row 0, cols 4–10: green bg, merge horizontally)
    for ci in range(CW_START, Q_COL + 1):
        cell = table.cell(0, ci)
        _set_cell_shading(cell, _CLR["HDR_GREEN"])
        _set_cell_text(cell, "", size=7)
    _set_hmerge_span(table, 0, CW_START, Q_COL)

    # Right text headers (dark blue, vertically merged)
    for i, label in enumerate(hdr_right):
        ci = Q_COL + 1 + i
        cell = table.cell(0, ci)
        _set_cell_shading(cell, _CLR["HDR_BLUE"])
        _set_cell_text(cell, label, size=7, bold=True, color=_CLR["WHITE"])
        _set_vmerge_restart(cell)
        cell1 = table.cell(1, ci)
        _set_cell_shading(cell1, _CLR["HDR_BLUE"])
        _set_cell_text(cell1, "", size=7)
        _set_vmerge_continue(cell1)

    # CW number cells (row 1, cols 4–9: green bg, white bold number)
    for i, cw_num in enumerate(cw_columns):
        ci = CW_START + i
        cell = table.cell(1, ci)
        _set_cell_shading(cell, _CLR["HDR_GREEN"])
        _set_cell_text(cell, str(cw_num), size=7, bold=True,
                       color=_CLR["WHITE"], align=WD_ALIGN_PARAGRAPH.CENTER)

    # Quarter label (row 1, col 10)
    q_cell = table.cell(1, Q_COL)
    _set_cell_shading(q_cell, _CLR["HDR_GREEN"])
    _set_cell_text(q_cell, quarter_label, size=7, bold=True,
                   color=_CLR["WHITE"], align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Data rows ────────────────────────────────────────────────────
    pg_ranges = _compute_product_group_ranges(all_rows)

    # Build a set of (data_row_idx → is_first_in_group, is_continuation)
    pg_first: set[int] = set()
    pg_cont:  set[int] = set()
    for start, end, _label in pg_ranges:
        pg_first.add(start)
        for idx in range(start + 1, end + 1):
            pg_cont.add(idx)

    for ri_data, row in enumerate(all_rows):
        ri = ri_data + 2   # table row index (offset by 2 header rows)

        # Product Group column — amber bg
        pg_cell = table.cell(ri, 0)
        _set_cell_shading(pg_cell, _CLR["AMBER"])
        if ri_data in pg_first:
            pg_text = row.get("product_group", "")
            _set_cell_text(pg_cell, pg_text, size=6, bold=True)
            # If this group spans multiple rows, mark as merge start
            if ri_data in {s for s, e, _ in pg_ranges if e > s}:
                _set_vmerge_restart(pg_cell)
        elif ri_data in pg_cont:
            _set_cell_text(pg_cell, "", size=6)
            _set_vmerge_continue(pg_cell)

        # Plant
        _set_cell_text(table.cell(ri, 1), row.get("plant", ""), size=6)

        # Customer / Channel
        _set_cell_text(table.cell(ri, 2), row.get("customer", ""), size=6)

        # KB Coverage (multi-line)
        _set_cell_text(table.cell(ri, 3), row.get("kb_coverage", ""), size=6)

        # CW coloured cells
        cw_colors: dict = row.get("cw_colors", {})
        for i, cw_num in enumerate(cw_columns):
            ci = CW_START + i
            rag = str(cw_colors.get(str(cw_num), "GREY")).upper()
            cell = table.cell(ri, ci)
            _set_cell_shading(cell, _RAG_HEX.get(rag, _CLR["GREY"]))
            _set_cell_text(cell, "", size=5)

        # Quarter column
        q_rag = str(row.get("quarter_color", "GREY")).upper()
        q_c = table.cell(ri, Q_COL)
        _set_cell_shading(q_c, _RAG_HEX.get(q_rag, _CLR["GREY"]))
        _set_cell_text(q_c, "", size=5)

        # Supplier
        _set_cell_text(table.cell(ri, 11), row.get("supplier", ""), size=6)

        # Comment
        _set_cell_text(table.cell(ri, 12), row.get("comment", ""), size=6)

        # FM Detail Letter
        _set_cell_text(table.cell(ri, 13), row.get("fm_detail", ""), size=6)

    # ── Compact row heights ──────────────────────────────────────────
    # python-docx doesn't expose row height directly in a convenient way,
    # but we can set it via XML to keep the table tight.
    for ri, row_obj in enumerate(table.rows):
        tr = row_obj._tr
        trPr = tr.get_or_add_trPr()
        # Remove any existing trHeight
        for existing in trPr.findall(qn("w:trHeight")):
            trPr.remove(existing)
        if ri < 2:
            # Header rows: slightly taller
            h = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="340" w:hRule="atLeast"/>')
        else:
            # Data rows: compact
            h = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="220" w:hRule="atLeast"/>')
        trPr.append(h)

    # ── Set thin borders on the whole table ──────────────────────────
    tbl_xml = table._tbl
    tblPr = tbl_xml.tblPr if tbl_xml.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>',
    )
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '</w:tblBorders>'
    )
    # Remove existing borders
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(parse_xml(borders_xml))

    # ── Footer: generation date ──────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_before = Pt(4)
    run = footer_para.add_run(f"Generated {date.today().isoformat()}")
    run.font.name = FONT
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor.from_string(_CLR["GREY"])

    # ── Save to buffer ───────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

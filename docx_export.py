"""
Markdown → DOCX Converter
==========================
Converts the agent's markdown output into a properly styled Word document.
Handles headings, bold, italic, bullet lists, code spans, and horizontal rules.
"""

import re
import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def _setup_styles(doc: Document):
    """Configure clean document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color_hex) in enumerate([
        (20, "2C5282"),   # Heading 1 — 20pt, blue
        (15, "2C5282"),   # Heading 2 — 15pt, blue
        (12, "1A1A1A"),   # Heading 3 — 12pt, dark
    ], start=1):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.size = Pt(size)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
        heading_style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        heading_style.paragraph_format.space_after = Pt(6)

    # Ensure List Bullet style exists and is configured
    try:
        list_style = doc.styles["List Bullet"]
    except KeyError:
        list_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
    list_style.font.name = "Calibri"
    list_style.font.size = Pt(11)
    list_style.paragraph_format.space_after = Pt(2)


def _parse_inline(paragraph, text: str):
    """Parse inline markdown (bold, italic, code, plain) and add runs to a paragraph."""
    # Pattern: **bold**, *italic*, `code`, or plain text
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)')
    pos = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

        pos = match.end()

    # Remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _detect_heading(line: str):
    """Return (level, text) if line is a heading, else None."""
    m = re.match(r'^(#{1,3})\s+(.+)$', line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return None


def _detect_bullet(line: str):
    """Return (indent_level, text) if line is a bullet, else None."""
    m = re.match(r'^(\s*)([-•*])\s+(.+)$', line)
    if m:
        indent = len(m.group(1))
        level = min(indent // 2, 2)  # 0, 1, or 2
        return level, m.group(3).strip()
    return None


def _detect_hr(line: str) -> bool:
    """Check if line is a horizontal rule."""
    return bool(re.match(r'^[-*_]{3,}\s*$', line.strip()))


def markdown_to_docx(markdown_text: str, title: str = "") -> io.BytesIO:
    """
    Convert markdown text to a .docx file in memory.
    Returns a BytesIO buffer ready for streaming.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    _setup_styles(doc)

    # Optional document title
    if title:
        p = doc.add_heading(title, level=1)
        # Add a thin line under the title
        p.paragraph_format.space_after = Pt(12)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if _detect_hr(stripped):
            # Add a subtle spacer paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a bottom border via a thin line run
            run = p.add_run("─" * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Headings
        heading = _detect_heading(stripped)
        if heading:
            level, text = heading
            p = doc.add_heading(level=level)
            _parse_inline(p, text)
            i += 1
            continue

        # Bullet points
        bullet = _detect_bullet(stripped)
        if bullet:
            indent_level, text = bullet
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            _parse_inline(p, text)
            i += 1
            continue

        # Regular paragraph — collect consecutive non-special lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (not next_line or _detect_heading(next_line) or
                    _detect_bullet(next_line) or _detect_hr(next_line)):
                break
            para_lines.append(next_line)
            i += 1

        p = doc.add_paragraph()
        _parse_inline(p, " ".join(para_lines))

    # Write to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
